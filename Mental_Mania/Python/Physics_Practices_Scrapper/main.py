"""
Problem: Required H2 A-level Physics practices for tuition, and no longer had access to digital school notes
Objective: Discovered this freemium site (savemyexams), requiring subscription to access questions and answers in PDF format. Google search for these returned outdated content.
           Wanted to compile topical PDFs of all these valuable questions and answers which also conveniently range in order of difficulty (Easy -> Medium -> Hard)
Approach: The questions are all free, but only the 1st question has a free accompanying answer.
          Dragging the image of question and answer and opening in a new tab reveals they are all stored in a CDN
          However, the link is not simple and navigating via directories in link is futile and time-consuming
          Further inspection of the elements on the site reveals the links to the images of these questions and answers are all written in HTML.
          Simple BeautifulSoup scraped these sites and downloaded these images locally, appending them first in .docx word documents, then converting to a pdf.
"""
from bs4 import BeautifulSoup
import requests
from docx import Document
from docx.shared import Inches
import docx2pdf
import os

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/108.0.0.0 Safari/537.36",
    "Accept-Language": "en-US,en;q=0.5",
    "Accept-Encoding": "gzip, deflate",
}
# To change, the name of saved pdf
TOPIC = "20C. Nuclear"
QNS_SAVE_FILEPATH = f"/Users/.../Desktop/{TOPIC}_Structured_Q.docx"
ANSWERS_SAVE_FILEPATH = f"/Users/.../Desktop/{TOPIC}_Structured_A.docx"
QNS = Document()
QNS.add_heading(f"TOPIC: {TOPIC} (Structured)", level=5)
ANSWERS = Document()
ANSWERS.add_heading(f"TOPIC: {TOPIC} (Structured)", level=5)
count = 0
for difficulty in ["Easy", "Medium", "Hard"]:
    # To change, any subsite that uses similar format (MCQ/Structured)
    url = f"https://www.savemyexams.co.uk/a-level/physics/aqa/17/topic-questions/8-nuclear-physics/8-4-nuclear-fusion--fission/-/structured-questions/{difficulty.lower()}/"
    response = requests.get(url=url, headers=HEADERS)
    soup = BeautifulSoup(response.text, "lxml")
    pair = soup.select(selector=".resource-page .question-part")
    QNS.add_heading(difficulty, level=5)
    ANSWERS.add_heading(difficulty, level=5)
    for item in pair:
        qn = item.find("img").get("src")
        try:
            ans = item.find("script").text.split('src=\\"')[1].split('\\" srcset')[0]
        except IndexError:
            print("################### Using other .find()")
            ans = item.find("script").text.split("src='")[1].split("' srcset")[0]
        # Some sites like https://www.savemyexams.co.uk/a-level/physics/aqa/17/topic-questions/5-electricity/5-3-circuits--the-potential-divider/-/multiple-choice/easy/
        # are broken, the answer is written directly in HTML and not as a link to the image of the answer. These return errors and are ignored.
        try:
            qn_save_filepath = f"/Users/.../Desktop/test/qns/{difficulty}/{qn.split('uploads/')[1].split('/')[2].split('_MCQ')[0]}.jpg"
            ans_save_filepath = f"/Users/.../Desktop/test/answers/{difficulty}/{ans.split('uploads/')[1].split('/')[2].split('_MCQ')[0]}.jpg"
        except IndexError:
            print(f"#################################### BROKE for {qn}      {ans}")
            continue
        print(f"QN: {qn}")
        print(f"ANS: {ans}")
        with open(qn_save_filepath, 'wb') as handler:
            handler.write(requests.get(qn).content)
        with open(ans_save_filepath, 'wb') as handler:
            handler.write(requests.get(ans).content)
        count += 1
        QNS.add_heading(str(count), level=3)
        # Inches to specify image size. otherwise, too big
        QNS.add_picture(qn_save_filepath, width=Inches(5))
        ANSWERS.add_heading(str(count), level=3)
        ANSWERS.add_picture(ans_save_filepath, width=Inches(5))

QNS.save(QNS_SAVE_FILEPATH)
ANSWERS.save(ANSWERS_SAVE_FILEPATH)
# Converting .docx to .pdf
docx2pdf.convert(QNS_SAVE_FILEPATH, QNS_SAVE_FILEPATH.replace(".docx", ".pdf"))
docx2pdf.convert(ANSWERS_SAVE_FILEPATH, ANSWERS_SAVE_FILEPATH.replace(".docx", ".pdf"))
# Deleting unnecessary .docx files
os.remove(QNS_SAVE_FILEPATH)
os.remove(ANSWERS_SAVE_FILEPATH)